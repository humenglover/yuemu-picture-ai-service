import hashlib
import os
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader
)
from langchain_core.documents import Document
from typing import List

# 导入日志功能
from utils.log_utils import knowledge_logger


def log_event(event_type: str, message: str, details: dict = None):
    """记录日志事件"""
    try:
        log_message = '[' + event_type + '] ' + message

        if details:
            log_message += ' | 详情: ' + str(details)

        knowledge_logger.info(log_message)
    except Exception as e:
        knowledge_logger.error('Logging error: ' + str(e))


def calculate_md5(file_path: str) -> str:
    """计算文件MD5"""
    hash_md5 = hashlib.md5()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(4096), b""):
            hash_md5.update(chunk)
    return hash_md5.hexdigest()


def load_documents_from_file(file_path: str) -> List[Document]:
    """从文件加载文档 - 简化版本，优先使用基本加载器"""
    file_ext = os.path.splitext(file_path)[1].lower()

    # 优先使用最稳定的加载器
    if file_ext == '.txt':
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            return loader.load()
        except Exception as e:
            knowledge_logger.error(f"TextLoader失败: {str(e)}")
            raise

    elif file_ext == '.pdf':
        try:
            loader = PyPDFLoader(file_path)
            return loader.load()
        except Exception as e:
            knowledge_logger.error(f"PyPDFLoader失败: {str(e)}")
            raise

    elif file_ext == '.docx':
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # 同样提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
                            
            text_content = '\n'.join(full_text)
            return [Document(page_content=text_content, metadata={"source": file_path})]
        except Exception as e:
            knowledge_logger.error(f"解析DOCX文件失败 {file_path}: {str(e)}")
            raise Exception(f"无法解析DOCX文件 {file_path}: {str(e)}")

    elif file_ext == '.md':
        # Markdown文件优先使用基本文本加载器
        try:
            loader = TextLoader(file_path, encoding='utf-8')
            return loader.load()
        except Exception as e:
            knowledge_logger.error(f"无法解析Markdown文件 {file_path}: {str(e)}")
            raise Exception(f"无法解析Markdown文件 {file_path}: {str(e)}")

    else:
        # 不支持的文件类型
        raise ValueError("Unsupported file type: " + file_ext)


def check_md5_hex(md5_for_check: str, md5_hex_store_path: str) -> bool:
    """检查MD5是否已存在"""
    # 确保目录存在
    os.makedirs(os.path.dirname(md5_hex_store_path), exist_ok=True)

    if not os.path.exists(md5_hex_store_path):
        # 创建文件
        with open(md5_hex_store_path, "w", encoding="utf-8") as f:
            pass
        return False

    with open(md5_hex_store_path, "r", encoding="utf-8") as f:
        for line in f.readlines():
            line = line.strip()
            if line == md5_for_check:
                return True

    return False


def save_md5_hex(md5_for_check: str, md5_hex_store_path: str):
    """保存MD5到记录文件"""
    # 确保目录存在
    os.makedirs(os.path.dirname(md5_hex_store_path), exist_ok=True)

    with open(md5_hex_store_path, "a", encoding="utf-8") as f:
        f.write(md5_for_check + "\n")


def get_file_documents(read_path: str) -> List[Document]:
    """根据文件扩展名获取文档加载器 - 简化稳定版本"""
    try:
        # 检查文件是否存在且不为空
        if not os.path.exists(read_path):
            knowledge_logger.warning(f"文件不存在: {read_path}")
            return []

        if os.path.getsize(read_path) == 0:
            knowledge_logger.warning(f"文件为空: {read_path}")
            return []

        file_ext = os.path.splitext(read_path)[1].lower()

        # 优先使用最稳定的加载器
        if file_ext == '.txt':
            loader = TextLoader(read_path, encoding='utf-8')
            return loader.load()
        elif file_ext == '.pdf':
            loader = PyPDFLoader(read_path)
            return loader.load()
        elif file_ext == '.docx':
            try:
                import docx
                doc = docx.Document(read_path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text.strip())
                text_content = '\n'.join(full_text)
                return [Document(page_content=text_content, metadata={"source": read_path})]
            except Exception as e:
                knowledge_logger.error(f"解析DOCX文件失败: {str(e)}")
                return []
        elif file_ext == '.md':
            # Markdown优先使用基本文本加载器
            loader = TextLoader(read_path, encoding='utf-8')
            return loader.load()
        else:
            knowledge_logger.warning(f"不支持的文件类型: {read_path}")
            return []

    except Exception as e:
        knowledge_logger.error(f"加载文件文档失败 | path: {read_path} | error: {str(e)}")
        return []


def find_filename_by_md5(target_md5: str, md5_store_path: str) -> str:
    """根据MD5值查找已存在的文件名"""
    if not os.path.exists(md5_store_path):
        return None

    with open(md5_store_path, "r", encoding="utf-8") as f:
        for line in f.readlines():
            line = line.strip()
            if line.startswith(target_md5 + ":"):
                # 如果记录格式是 md5:filename
                parts = line.split(':', 1)
                if len(parts) > 1:
                    return parts[1]
            elif line == target_md5:
                # 如果记录格式只有MD5值，返回None
                return None
    return None


def save_md5_with_filename(md5_for_check: str, filename: str, md5_hex_store_path: str):
    """保存MD5和文件名到记录文件"""
    # 确保目录存在
    os.makedirs(os.path.dirname(md5_hex_store_path), exist_ok=True)

    with open(md5_hex_store_path, "a", encoding="utf-8") as f:
        f.write(md5_for_check + ":" + filename + "\n")


def check_md5_exists_only(md5_for_check: str, md5_hex_store_path: str) -> bool:
    """仅检查MD5是否已存在（不考虑文件名）"""
    # 确保目录存在
    os.makedirs(os.path.dirname(md5_hex_store_path), exist_ok=True)

    if not os.path.exists(md5_hex_store_path):
        # 创建文件
        with open(md5_hex_store_path, "w", encoding="utf-8") as f:
            pass
        return False

    with open(md5_hex_store_path, "r", encoding="utf-8") as f:
        for line in f.readlines():
            line = line.strip()
            # 提取MD5部分进行比较
            md5_part = line.split(':')[0] if ':' in line else line
            if md5_part == md5_for_check:
                return True

    return False


def check_md5_exists_permanently(md5_for_check: str, permanent_md5_store_path: str) -> bool:
    """检查MD5是否曾经被处理过（永久记录），用于严格的内容重复检测"""
    os.makedirs(os.path.dirname(permanent_md5_store_path), exist_ok=True)

    if not os.path.exists(permanent_md5_store_path):
        return False

    with open(permanent_md5_store_path, 'r', encoding='utf-8') as f:
        for line in f.readlines():
            line = line.strip()
            if line == md5_for_check:
                return True

    return False


def save_md5_permanently(md5_for_check: str, permanent_md5_store_path: str):
    """永久保存MD5记录，用于严格的内容重复检测"""
    os.makedirs(os.path.dirname(permanent_md5_store_path), exist_ok=True)

    # 先检查是否已存在
    if check_md5_exists_permanently(md5_for_check, permanent_md5_store_path):
        return

    with open(permanent_md5_store_path, 'a', encoding='utf-8') as f:
        f.write(md5_for_check + '\n')